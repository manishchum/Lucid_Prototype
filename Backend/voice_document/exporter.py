from docx import Document
from openpyxl import Workbook
import tempfile


def generate_docx(renderable_content):
    doc = Document()

    doc.add_heading(
        renderable_content.get("title", "Voice Document"),
        level=1
    )

    for section in renderable_content.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=2)
        doc.add_paragraph(section.get("body", ""))

        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    if renderable_content.get("action_items"):
        doc.add_heading("Action Items", level=2)

        for item in renderable_content["action_items"]:
            doc.add_paragraph(
                f"{item.get('owner')} - {item.get('task')} "
                f"(Due: {item.get('due_date')})"
            )

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    )

    doc.save(temp_file.name)

    return temp_file.name


def generate_xlsx(renderable_content):
    wb = Workbook()

    sheets = renderable_content.get("sheets", [])

    first_sheet = True

    for sheet_data in sheets:

        if first_sheet:
            ws = wb.active
            ws.title = sheet_data["name"]
            first_sheet = False
        else:
            ws = wb.create_sheet(sheet_data["name"])

        ws.append(sheet_data["columns"])

        for row in sheet_data["rows"]:
            ws.append(row)

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".xlsx"
    )

    wb.save(temp_file.name)

    return temp_file.name
